from __future__ import annotations

import io
import json
import subprocess
import tempfile
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Any

from docx.image.exceptions import UnrecognizedImageError
from docx.image.image import Image
from docx.shared import Inches
from fastapi import HTTPException
from openpyxl import load_workbook

from services.datamerge import docx


THIS_DIR = Path(__file__).resolve().parent
SERVER_ROOT = THIS_DIR.parents[1]
DEFAULT_TCD09_IMAGE_RULES_PATH = SERVER_ROOT / "config" / "tcd09_image_rules.json"

DRAWING_NAMESPACES = {
    "main": "http://schemas.openxmlformats.org/spreadsheetml/2006/main",
    "rel": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "xdr": "http://schemas.openxmlformats.org/drawingml/2006/spreadsheetDrawing",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
}


def _load_json(path: str | Path, label: str) -> dict[str, Any]:
    json_path = Path(path)
    if not json_path.is_file():
        raise HTTPException(status_code=500, detail=f"{label} does not exist: {json_path}")

    try:
        with json_path.open("r", encoding="utf-8-sig") as file:
            data = json.load(file)
    except json.JSONDecodeError as exc:
        raise HTTPException(
            status_code=500,
            detail=(
                f"{label} is invalid JSON: {json_path}. "
                f"Line {exc.lineno}, column {exc.colno}: {exc.msg}"
            ),
        ) from exc
    except OSError as exc:
        raise HTTPException(status_code=500, detail=f"Unable to read {label}: {json_path}") from exc

    if not isinstance(data, dict):
        raise HTTPException(status_code=500, detail=f"{label} root must be a JSON object.")
    return data


def load_tcd09_image_rules(
    config_path: str | Path = DEFAULT_TCD09_IMAGE_RULES_PATH,
) -> dict[str, Any]:
    config = _load_json(config_path, "TCD09 image rules")
    sections = config.get("sections")
    if not isinstance(sections, list) or not sections:
        raise HTTPException(status_code=500, detail='"sections" must be a non-empty JSON array.')

    for index, section in enumerate(sections, start=1):
        if not isinstance(section, dict):
            raise HTTPException(status_code=500, detail=f"Section rule #{index} must be an object.")

        for required_key in ["key", "placeholder", "sheet_name", "start_label"]:
            value = section.get(required_key)
            if not isinstance(value, str) or not value.strip():
                raise HTTPException(
                    status_code=500,
                    detail=f'Section rule #{index} must define non-empty string "{required_key}".',
                )

        # end_label is optional. When present, it must identify the first row
        # that does NOT belong to the current image section.
        end_label = section.get("end_label")
        if end_label is not None and (
            not isinstance(end_label, str) or not end_label.strip()
        ):
            raise HTTPException(
                status_code=500,
                detail=(
                    f'Section rule #{index} has invalid "end_label". '
                    "Omit it for the final section, or provide a non-empty string."
                ),
            )

        customer_start_column = section.get("customer_start_column", 9)
        if not isinstance(customer_start_column, int) or customer_start_column < 1:
            raise HTTPException(
                status_code=500,
                detail=f'Section rule #{index} has invalid "customer_start_column".',
            )

        image_width_inches = section.get("image_width_inches", 2.35)
        if not isinstance(image_width_inches, (int, float)) or image_width_inches <= 0:
            raise HTTPException(
                status_code=500,
                detail=f'Section rule #{index} has invalid "image_width_inches".',
            )

    return config


def _normalize_text(value: Any) -> str:
    return str(value or "").replace("\n", " ").replace("\r", " ").strip()


def _canonicalize_title(value: Any) -> str:
    return (
        _normalize_text(value)
        .replace("（", "(")
        .replace("）", ")")
        .replace(" ", "")
        .casefold()
    )


def _find_label_cell(worksheet, label: str, customer_start_column: int) -> tuple[int, int]:
    expected = _canonicalize_title(label)
    for row in worksheet.iter_rows():
        for cell in row:
            if int(cell.column) < customer_start_column:
                continue
            if _canonicalize_title(cell.value) == expected:
                return int(cell.row), int(cell.column)

    raise HTTPException(
        status_code=400,
        detail=f'Label not found on customer side: sheet={worksheet.title} label="{label}"',
    )


def _find_label_cell_after(
    worksheet,
    label: str,
    customer_start_column: int,
    *,
    min_row: int,
) -> tuple[int, int]:
    """Find an exact label at or below ``min_row`` in the customer area."""

    expected = _canonicalize_title(label)
    for row in worksheet.iter_rows(
        min_row=max(1, int(min_row)),
        max_row=int(worksheet.max_row),
        min_col=int(customer_start_column),
        max_col=int(worksheet.max_column),
    ):
        for cell in row:
            if _canonicalize_title(cell.value) == expected:
                return int(cell.row), int(cell.column)

    raise HTTPException(
        status_code=400,
        detail=(
            f'End label not found after the start section: '
            f'sheet={worksheet.title} label="{label}" min_row={min_row}'
        ),
    )


def _resolve_section_window(
    xlsm_bytes: bytes,
    *,
    sheet_name: str,
    start_label: str,
    customer_start_column: int,
    end_label: str | None = None,
) -> dict[str, int | str]:
    """Resolve an image section using explicit start/end labels.

    ``start_label`` is included in the section. ``end_label`` is excluded and
    therefore marks the first row of the next section. When ``end_label`` is
    omitted, the current section is treated as the final section and extends to
    the end of the worksheet.
    """

    workbook = load_workbook(io.BytesIO(xlsm_bytes), data_only=False, keep_vba=True)
    try:
        try:
            sheet = workbook[sheet_name]
        except KeyError as exc:
            raise HTTPException(
                status_code=400,
                detail=f"Sheet not found in Excel template: {sheet_name}",
            ) from exc

        start_row, _ = _find_label_cell(
            sheet,
            start_label,
            customer_start_column,
        )

        normalized_end_label = str(end_label or "").strip()
        if normalized_end_label:
            end_row, _ = _find_label_cell_after(
                sheet,
                normalized_end_label,
                customer_start_column,
                min_row=start_row + 1,
            )
        else:
            end_row = int(sheet.max_row) + 1

        if end_row <= start_row:
            raise HTTPException(
                status_code=400,
                detail=(
                    f"Invalid section window for sheet={sheet_name}. "
                    f"start_label={start_label!r}, end_label={normalized_end_label!r}, "
                    f"start_row={start_row}, end_row={end_row}"
                ),
            )

        return {
            "sheet_name": sheet_name,
            "sheet_index": workbook.sheetnames.index(sheet_name) + 1,
            "start_row": start_row,
            "end_row": end_row,
            "customer_start_column": customer_start_column,
        }
    finally:
        workbook.close()

def _is_supported_docx_image(image_path: Path) -> bool:
    try:
        Image.from_file(str(image_path))
        return True
    except (FileNotFoundError, UnrecognizedImageError, OSError):
        return False


def _convert_emf_to_png(source_path: Path, output_path: Path) -> Path:
    src = str(source_path).replace("'", "''")
    dst = str(output_path).replace("'", "''")
    ps_script = (
        "Add-Type -AssemblyName System.Drawing; "
        f"$src='{src}'; "
        f"$dst='{dst}'; "
        "$img=[System.Drawing.Image]::FromFile($src); "
        "$bmp=New-Object System.Drawing.Bitmap $img.Width, $img.Height; "
        "$graphics=[System.Drawing.Graphics]::FromImage($bmp); "
        "$graphics.Clear([System.Drawing.Color]::White); "
        "$graphics.DrawImage($img, 0, 0, $img.Width, $img.Height); "
        "$bmp.Save($dst, [System.Drawing.Imaging.ImageFormat]::Png); "
        "$graphics.Dispose(); $bmp.Dispose(); $img.Dispose()"
    )
    result = subprocess.run(
        ["powershell", "-NoProfile", "-Command", ps_script],
        capture_output=True,
        text=True,
        check=False,
    )
    if result.returncode != 0 or not output_path.exists():
        stderr = result.stderr.strip() or result.stdout.strip() or "unknown PowerShell error"
        raise HTTPException(status_code=500, detail=f"Failed to convert EMF to PNG: {stderr}")
    return output_path


def _materialize_docx_image(raw_bytes: bytes, media_name: str, temp_dir: Path) -> Path:
    source_path = temp_dir / Path(media_name).name
    source_path.write_bytes(raw_bytes)

    if _is_supported_docx_image(source_path):
        return source_path

    if source_path.suffix.lower() == ".emf":
        converted_path = _convert_emf_to_png(source_path, source_path.with_suffix(".png"))
        if _is_supported_docx_image(converted_path):
            return converted_path

    raise HTTPException(
        status_code=400,
        detail=f"Unsupported image format for Word insertion: {source_path.name}",
    )



def _get_flow_image_width(image_path: Path, max_width_inches: float):
    """Return an image width suitable for natural inline Word layout.

    Small images keep their intrinsic width. Oversized images are reduced to
    the existing ``image_width_inches`` limit. Height is omitted so python-docx
    preserves the original aspect ratio.
    """

    max_width = Inches(float(max_width_inches))
    try:
        image = Image.from_file(str(image_path))
        natural_width = image.width
    except (FileNotFoundError, UnrecognizedImageError, OSError):
        return max_width

    if natural_width is None or int(natural_width) <= 0:
        return max_width

    return min(natural_width, max_width)


def _iter_images_for_window(
    xlsm_bytes: bytes,
    window: dict[str, int | str],
) -> list[tuple[int, int, str, bytes]]:
    sheet_index = int(window["sheet_index"])
    start_row = int(window["start_row"])
    end_row = int(window["end_row"])
    customer_start_column = int(window["customer_start_column"])

    with zipfile.ZipFile(io.BytesIO(xlsm_bytes)) as workbook_zip:
        workbook_xml = ET.fromstring(workbook_zip.read("xl/workbook.xml"))
        sheets = workbook_xml.find("main:sheets", DRAWING_NAMESPACES)
        if sheets is None or len(list(sheets)) < sheet_index:
            raise HTTPException(
                status_code=400,
                detail="Configured sheet index is missing in Excel template.",
            )

        target_sheet = list(sheets)[sheet_index - 1]
        target_rid = target_sheet.attrib.get(f'{{{DRAWING_NAMESPACES["rel"]}}}id')
        if not target_rid:
            raise HTTPException(
                status_code=400,
                detail="Sheet relationship id missing in Excel template.",
            )

        workbook_rels = ET.fromstring(workbook_zip.read("xl/_rels/workbook.xml.rels"))
        sheet_target = None
        for rel in workbook_rels:
            if rel.attrib.get("Id") == target_rid:
                sheet_target = rel.attrib.get("Target")
                break
        if not sheet_target:
            raise HTTPException(
                status_code=400,
                detail="Sheet target missing in Excel template.",
            )

        sheet_xml = ET.fromstring(workbook_zip.read("xl/" + sheet_target.lstrip("/")))
        drawing = sheet_xml.find("main:drawing", DRAWING_NAMESPACES)
        if drawing is None:
            return []

        drawing_rid = drawing.attrib.get(f'{{{DRAWING_NAMESPACES["rel"]}}}id')
        sheet_rels_path = "xl/worksheets/_rels/" + Path(sheet_target).name + ".rels"
        sheet_rels = ET.fromstring(workbook_zip.read(sheet_rels_path))

        drawing_target = None
        for rel in sheet_rels:
            if rel.attrib.get("Id") == drawing_rid:
                drawing_target = rel.attrib.get("Target")
                break
        if not drawing_target:
            return []

        drawing_xml = ET.fromstring(
            workbook_zip.read("xl/drawings/" + Path(drawing_target).name)
        )
        drawing_rels_path = (
            "xl/drawings/_rels/" + Path(drawing_target).name + ".rels"
        )
        drawing_rels = ET.fromstring(workbook_zip.read(drawing_rels_path))
        rel_map = {
            rel.attrib.get("Id"): rel.attrib.get("Target")
            for rel in drawing_rels
        }

        images: list[tuple[int, int, str, bytes]] = []
        for anchor in list(drawing_xml):
            from_node = anchor.find("xdr:from", DRAWING_NAMESPACES)
            pic = anchor.find("xdr:pic", DRAWING_NAMESPACES)
            if from_node is None or pic is None:
                continue

            row_node = from_node.find("xdr:row", DRAWING_NAMESPACES)
            col_node = from_node.find("xdr:col", DRAWING_NAMESPACES)
            if row_node is None or col_node is None:
                continue

            row = int(row_node.text) + 1
            col = int(col_node.text) + 1
            if not (
                start_row <= row < end_row
                and col >= customer_start_column
            ):
                continue

            blip = pic.find(".//a:blip", DRAWING_NAMESPACES)
            if blip is None:
                continue

            rel_id = blip.attrib.get(f'{{{DRAWING_NAMESPACES["r"]}}}embed')
            media_target = rel_map.get(rel_id)
            if not media_target:
                continue

            media_path = "xl/media/" + Path(media_target).name
            images.append((row, col, media_path, workbook_zip.read(media_path)))

        images.sort(key=lambda item: (item[0], item[1], item[2]))
        return images


def insert_excel_section_images(
    document_stream: io.BytesIO,
    xlsm_bytes: bytes,
    config_path: str | Path = DEFAULT_TCD09_IMAGE_RULES_PATH,
) -> io.BytesIO:
    config = load_tcd09_image_rules(config_path)
    sections = config.get("sections", [])

    document_stream.seek(0)
    document = docx.Document(document_stream)

    with tempfile.TemporaryDirectory(prefix="puma_tcd09_images_") as temp_root:
        temp_dir = Path(temp_root)
        replaced_any = False

        for section in sections:
            placeholder = str(section["placeholder"]).strip()
            placeholder_paragraph = None
            for paragraph in document.paragraphs:
                if str(getattr(paragraph, "text", "")).strip() == placeholder:
                    placeholder_paragraph = paragraph
                    break

            if placeholder_paragraph is None:
                continue

            sheet_name = str(section["sheet_name"]).strip()
            end_label_value = section.get("end_label")
            end_label = (
                str(end_label_value).strip()
                if isinstance(end_label_value, str) and end_label_value.strip()
                else None
            )
            window = _resolve_section_window(
                xlsm_bytes,
                sheet_name=sheet_name,
                start_label=str(section["start_label"]).strip(),
                customer_start_column=int(section.get("customer_start_column", 9)),
                end_label=end_label,
            )
            images = _iter_images_for_window(xlsm_bytes, window)
            if not images:
                continue

            placeholder_paragraph.clear()
            max_image_width_inches = float(section.get("image_width_inches", 2.35))

            for index, (_, _, media_name, raw_bytes) in enumerate(images):
                image_path = _materialize_docx_image(raw_bytes, media_name, temp_dir)
                image_width = _get_flow_image_width(
                    image_path,
                    max_width_inches=max_image_width_inches,
                )

                # Keep every picture as an inline object in the same paragraph.
                # The regular spaces between pictures allow Word to wrap them
                # naturally according to the remaining line width.
                picture_run = placeholder_paragraph.add_run()
                picture_run.add_picture(str(image_path), width=image_width)

                if index < len(images) - 1:
                    placeholder_paragraph.add_run("  ")

            replaced_any = True

    if not replaced_any:
        document_stream.seek(0)
        return document_stream

    output_stream = io.BytesIO()
    document.save(output_stream)
    output_stream.seek(0)
    return output_stream
